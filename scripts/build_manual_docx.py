#!/usr/bin/env python3
"""
Build a formatted product manual DOCX using python-docx and Pandoc.
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PANDOC_PATH = Path(r"C:\Program Files\Pandoc\pandoc.exe")


def safe_path_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*\x00-\x1f]+', "-", value.strip())
    cleaned = re.sub(r"\s+", "-", cleaned)
    cleaned = re.sub(r"-{2,}", "-", cleaned).strip(" .-")
    return cleaned or "system-app"


def display_system_name(system_name: str) -> str:
    name = system_name.strip()
    return name if name.endswith("系统") else f"{name}系统"


def derive_short_name(system_name: str) -> str:
    name = display_system_name(system_name)
    replacements = [
        ("高校学生信息管理系统", "高校管理系统"),
        ("学生信息管理系统", "学生管理系统"),
        ("信息管理系统", "管理系统"),
        ("业务管理系统", "管理系统"),
        ("服务管理系统", "管理系统"),
        ("平台管理系统", "管理系统"),
    ]
    for source, target in replacements:
        if source in name:
            shortened = name.replace(source, target)
            if shortened != name:
                return shortened
    return name


def load_manifest(root: Path, system_name: str) -> dict:
    direct_manifest = root / safe_path_name(system_name) / "docs" / "Template" / "delivery-manifest.json"
    if direct_manifest.exists():
        return json.loads(direct_manifest.read_text(encoding="utf-8"))
    return {}


def read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def clean_version(raw: str) -> str:
    value = str(raw).strip()
    value = value.lstrip("^~<>= ")
    return value or str(raw)


def normalize_screenshot_key(path: Path) -> str:
    stem = path.stem.lower().replace("_", "-")
    stem = re.sub(r"^\d+[-\s]*", "", stem)
    stem = re.sub(r"-{2,}", "-", stem).strip("- ")
    return stem


def label_from_filename(path: Path) -> str:
    key = normalize_screenshot_key(path)
    whole_map = {
        "login": "登录页面",
        "signin": "登录页面",
        "sign-in": "登录页面",
        "dashboard": "控制台页面",
        "overview": "总览页面",
        "home": "首页页面",
        "student-detail": "学生详情页面",
        "product-detail": "商品详情页面",
        "order-detail": "订单详情页面",
        "user-detail": "用户详情页面",
    }
    token_map = {
        "student": "学生",
        "students": "学生档案",
        "teacher": "教师",
        "teachers": "教师",
        "course": "课程",
        "courses": "课程",
        "class": "班级",
        "classes": "班级",
        "grade": "成绩",
        "grades": "成绩",
        "attendance": "考勤",
        "dashboard": "控制台",
        "overview": "总览",
        "home": "首页",
        "report": "报表",
        "reports": "报表",
        "setting": "系统设置",
        "settings": "系统设置",
        "user": "用户",
        "users": "用户管理",
        "role": "角色",
        "roles": "角色权限",
        "permission": "权限",
        "permissions": "权限",
        "workflow": "流程",
        "approval": "审批",
        "product": "商品",
        "products": "商品",
        "inventory": "库存",
        "order": "订单",
        "orders": "订单",
        "customer": "客户",
        "customers": "客户",
        "finance": "财务",
        "payment": "支付",
        "monitor": "监控",
        "audit": "审计",
        "message": "消息",
        "messages": "消息",
        "notification": "通知",
        "notifications": "通知",
        "profile": "资料",
        "detail": "详情",
        "details": "详情",
        "list": "列表",
        "form": "表单",
        "create": "新建",
        "edit": "编辑",
        "analytics": "分析",
        "analysis": "分析",
        "search": "检索",
        "schedule": "日程",
        "exam": "考试",
        "library": "图书馆",
        "resource": "资源",
        "resources": "资源",
    }
    if key in whole_map:
        return whole_map[key]
    if re.search(r"[\u4e00-\u9fff]", path.stem):
        return path.stem if path.stem.endswith("页面") else f"{path.stem}页面"
    parts = [token_map.get(part) for part in key.split("-") if part]
    parts = [part for part in parts if part]
    if not parts:
        return "功能页面"
    title = "".join(parts)
    return title if title.endswith("页面") else f"{title}页面"


def infer_modules(system_dir: Path) -> list[str]:
    modules: list[str] = []
    seen: set[str] = set()

    module_plan_file = system_dir / "docs" / "Template" / "module-plan.md"
    if module_plan_file.exists():
        text = module_plan_file.read_text(encoding="utf-8", errors="ignore")
        for line in text.splitlines():
            match = re.match(r"^\s*(?:\d+[\.\、]\s*|[-*]\s+)(.+?)\s*$", line)
            if match:
                name = match.group(1).strip()
                if name and name not in seen:
                    seen.add(name)
                    modules.append(name)

    frontend_modules_dir = system_dir / "code" / "frontend" / "src" / "modules"
    if frontend_modules_dir.exists():
        for child in sorted(frontend_modules_dir.iterdir()):
            if child.is_dir():
                name = child.name.strip()
                if name and name not in seen:
                    seen.add(name)
                    modules.append(name)

    backend_router_dir = system_dir / "code" / "backend" / "app" / "routers"
    if backend_router_dir.exists():
        for child in sorted(backend_router_dir.glob("*.py")):
            if child.stem == "__init__":
                continue
            name = child.stem.strip()
            if name and name not in seen:
                seen.add(name)
                modules.append(name)

    if not modules:
        modules = ["仪表盘", "基础资料", "核心业务管理", "流程审批", "统计报表", "系统设置"]

    return modules[:10]


def infer_frontend_language(system_dir: Path) -> str:
    candidates = [
        system_dir / "code" / "frontend" / "package.json",
        system_dir / "demo" / "package.json",
    ]
    for candidate in candidates:
        package_data = read_json(candidate)
        if package_data:
            deps = package_data.get("dependencies", {})
            dev_deps = package_data.get("devDependencies", {})
            ts_version = dev_deps.get("typescript") or deps.get("typescript")
            if ts_version:
                return f"TypeScript {clean_version(ts_version)}"
            react_version = deps.get("react")
            if react_version:
                return f"JavaScript {clean_version(react_version)}"
    return "TypeScript 5.0"


def infer_backend_language(system_dir: Path) -> str:
    pyproject = system_dir / "code" / "backend" / "pyproject.toml"
    if pyproject.exists():
        text = pyproject.read_text(encoding="utf-8", errors="ignore")
        match = re.search(r"requires-python\s*=\s*['\"]([^'\"]+)['\"]", text)
        if match:
            version = match.group(1).replace(">=", "").strip()
            return f"Python {version}"

    backend_dir = system_dir / "code" / "backend"
    if backend_dir.exists():
        if list(backend_dir.rglob("*.py")):
            return "Python 3.10"
        if list(backend_dir.rglob("*.java")):
            return "Java 17"
        if list(backend_dir.rglob("*.ts")):
            return "TypeScript 5.0"

    return "Python 3.10"


def infer_fullstack_languages(system_dir: Path) -> str:
    frontend = infer_frontend_language(system_dir)
    backend = infer_backend_language(system_dir)
    if frontend == backend:
        return frontend
    return f"{frontend}，{backend}"


def load_manual_content(system_dir: Path) -> dict:
    template_dir = system_dir / "docs" / "Template"
    candidates = list(template_dir.glob("*-manual-content.json"))
    if not candidates:
        return {}
    try:
        return json.loads(candidates[0].read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def build_purpose_paragraph(system_name: str, modules: list[str]) -> str:
    module_sample = "、".join(modules[:4]) if modules else "核心业务"
    return (
        f"{system_name}面向业务管理、信息维护和流程协同等实际场景，围绕{module_sample}等模块构建统一工作平台。"
        "系统帮助使用人员通过标准化界面完成基础数据录入、业务处理、结果查询和过程跟踪，减少人工台账与重复沟通成本，"
        "提升信息准确性、处理效率和管理透明度，适用于规范化运营、综合管理、统计分析和业务协同等工作场景。"
    )


def build_function_paragraph(system_name: str, modules: list[str]) -> str:
    module_text = "、".join(modules[:8]) if modules else "首页总览、模块导航、列表查询、表单处理"
    return (
        f"{system_name}围绕{module_text}等功能模块进行设计，提供首页总览、数据录入、列表查询、详情查看、业务处理、统计分析和系统配置等能力。"
        "各模块在统一界面中协同工作，支持用户按角色完成信息维护、业务执行、状态跟踪和结果复核，并通过清晰的操作入口与结构化信息呈现提升整体使用效率。"
    )


def screenshot_paragraph(title: str) -> str:
    return (
        f"{title}页面用于展示当前模块的主要功能界面和核心业务内容，页面中通常包含关键数据区、"
        "业务入口、状态信息、筛选条件以及用户完成本模块任务所需的主要操作控件。通过该页面，"
        "使用人员可以快速理解本模块承担的业务用途、信息结构和操作路径，并结合按钮、列表、"
        "统计区或详情区域完成查询、录入、查看或处理等常用操作，是系统日常使用与演示说明中的重要页面。"
    )


def choose_text(preferred: str | None, fallback: str) -> str:
    if preferred and preferred.strip():
        return preferred.strip()
    return fallback


def extract_screenshot_overrides(manual_content: dict) -> dict[int, dict]:
    result: dict[int, dict] = {}
    for item in manual_content.get("screenshot_sections", []) or []:
        index = item.get("index")
        try:
            index = int(index)
        except (TypeError, ValueError):
            continue
        result[index] = item
    return result


def has_placeholder_text(value: str | None) -> bool:
    if not value or not value.strip():
        return True
    text = value.strip()
    placeholders = [
        "[请",
        "[在此",
        "页面用于展示当前模块的主要功能界面和核心业务内容",
        "面向业务管理、信息维护和流程协同等实际场景",
        "围绕",
    ]
    return any(marker in text for marker in placeholders)


def validate_manual_content(manual_content: dict, screenshots: list[Path]) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []

    if has_placeholder_text(manual_content.get("short_name")):
        errors.append("manual-content.json 缺少有效的 short_name，需先由当前 agent 填写简称。")

    if has_placeholder_text(manual_content.get("purpose_text")):
        errors.append("manual-content.json 缺少有效的 purpose_text，需先由当前 agent 填写“软件用途”文案。")

    if has_placeholder_text(manual_content.get("function_text")):
        errors.append("manual-content.json 缺少有效的 function_text，需先由当前 agent 填写“软件功能”文案。")

    if has_placeholder_text(manual_content.get("development_environment_text")):
        errors.append("manual-content.json 缺少有效的 development_environment_text，需先由当前 agent 填写“软件开发环境”文案。")

    screenshot_overrides = extract_screenshot_overrides(manual_content)
    for index, screenshot in enumerate(screenshots, start=1):
        item = screenshot_overrides.get(index)
        if not item:
            errors.append(f"manual-content.json 缺少截图 {screenshot.name} 的说明内容。")
            continue
        title = item.get("title", "")
        if has_placeholder_text(title):
            errors.append(f"截图 {screenshot.name} 的 title 仍是占位或空白，需先由当前 agent 填写中文页面名。")
        if has_placeholder_text(item.get("description_text")):
            errors.append(f"截图 {screenshot.name} 的 description_text 仍是占位或模板文案，需先由当前 agent 填写。")
        if not item.get("title_options"):
            errors.append(f"截图 {screenshot.name} 缺少可供当前 agent 参考的中文标题候选。")

    return errors, warnings


def set_run_songti(run, size_pt: float, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph(
    document: Document,
    text: str,
    *,
    size_pt: float = 12,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent_chars: int = 0,
    style: str | None = None,
) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = align
    if first_line_indent_chars:
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_chars * 12)
    run = paragraph.add_run(text)
    set_run_songti(run, size_pt, bold)


def add_heading_paragraph(document: Document, text: str, level: int) -> None:
    style_name = "Heading 1" if level == 1 else "Heading 2"
    add_paragraph(document, text, size_pt=16, bold=True, style=style_name)


def set_table_layout(
    table,
    widths_cm: list[float],
    *,
    align=WD_TABLE_ALIGNMENT.CENTER,
    row_height_cm: float = 0.9,
    text_align=WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    table.alignment = align
    table.autofit = False
    for row in table.rows:
        row.height = Cm(row_height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = text_align
                for run in paragraph.runs:
                    set_run_songti(run, 12, False)


def add_toc_field(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-2" \h \z \u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "右键更新目录"
    placeholder_run.append(placeholder_text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    paragraph._p.append(fld_begin)
    paragraph._p.append(instr_text)
    paragraph._p.append(fld_separate)
    paragraph._p.append(placeholder_run)
    paragraph._p.append(fld_end)


def add_revision_table(document: Document, author: str) -> None:
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["版本号", "生成日期", "作者", "修订内容"]
    values = ["V1.0", "", author, "初始版本"]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        table.cell(1, index).text = values[index]
    set_table_layout(
        table,
        [2.0, 3.0, 6.5, 3.5],
        align=WD_TABLE_ALIGNMENT.CENTER,
        row_height_cm=0.9,
        text_align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9D9D9")
        tc_pr.append(shd)


def add_two_col_table(document: Document, headers: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = headers[0]
    table.cell(0, 1).text = headers[1]
    for row_index, (left, right) in enumerate(rows, start=1):
        table.cell(row_index, 0).text = left
        table.cell(row_index, 1).text = right
    set_table_layout(
        table,
        [2.8, 12.2],
        align=WD_TABLE_ALIGNMENT.LEFT,
        row_height_cm=0.9,
        text_align=WD_ALIGN_PARAGRAPH.LEFT,
    )


def maybe_add_picture(document: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    if image_path.exists() and image_path.stat().st_size > 0:
        try:
            picture_paragraph = document.add_paragraph()
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = picture_paragraph.add_run()
            run.add_picture(str(image_path), width=Cm(width_cm))
        except Exception:
            add_paragraph(document, f"[图片无法自动嵌入：{image_path.name}]", align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_paragraph(document, f"[图片缺失：{image_path.name}]", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(document, caption, align=WD_ALIGN_PARAGRAPH.CENTER)


def render_markdown_draft(
    system_name: str,
    short_name: str,
    frontend_language: str,
    backend_language: str,
    fullstack_languages: str,
    modules: list[str],
    screenshots: list[Path],
    manual_content: dict,
) -> str:
    parts = [
        f"{system_name}[简称：{short_name}]V1.0",
        "",
        "产品说明书",
        "",
        "\\newpage",
        "",
        "## 目  录",
        "",
        "[右键更新目录]",
        "",
        "\\newpage",
        "",
        "## 修订记录",
        "",
        "| 版本号 | 生成日期 | 作者 | 修订内容 |",
        "| --- | --- | --- | --- |",
        "| V1.0 |  | 孔祥鑫 | 初始版本 |",
        "",
        "\\newpage",
        "",
        "## 一、软件介绍",
        "",
        f"软件名称：{system_name}",
        "",
        f"简称：{short_name}",
        "",
        "版本号：V1.0",
        "",
        "软件类别：应用软件",
        "",
        "著作权人：孔祥鑫",
        "",
        "\\newpage",
        "",
        "## 二、软件用途",
        "",
        choose_text(manual_content.get("purpose_text"), build_purpose_paragraph(system_name, modules)),
        "",
        "\\newpage",
        "",
        "## 三、软件功能",
        "",
        choose_text(manual_content.get("function_text"), build_function_paragraph(system_name, modules)),
        "",
        "\\newpage",
        "",
        "## 四、运行环境",
        "",
        "### 4.1 硬件要求",
        "",
        "| 类型 | 基本要求 |",
        "| --- | --- |",
        "| 服务器端 | CPU 8核1.60GHz，内存8G，硬盘剩余空间10G |",
        "",
        "### 4.2 软件环境",
        "",
        "| 名称 | 基本环境 |",
        "| --- | --- |",
        "| 操作系统 | Windows 10 64位 |",
        "| 数据库软件 | MySQL8.0 |",
        "| 开发软件 | Opencode，Claude Code，Codex，IntelliJ IDEA，Navicat 16 |",
        f"| 开发语言 | {fullstack_languages} |",
        "",
        "\\newpage",
        "",
        "### 4.3 软件开发环境",
        "",
        choose_text(
            manual_content.get("development_environment_text"),
            f"本软件分为前端页面和后端业务逻辑，其中前端页面使用{frontend_language}进行开发，后端业务逻辑使用{backend_language}进行开发，数据库使用MySQL8.0，数据库管理采用Navicat 16，整个系统使用IntelliJ IDEA环境进行开发。开发界面如图4-1所示。",
        ),
        "",
        "[在此处粘贴开发界面图片]",
        "",
        "图4-1 软件开发界面",
        "",
        "\\newpage",
        "",
        "## 五、软件使用",
        "",
    ]

    screenshot_overrides = extract_screenshot_overrides(manual_content)
    if screenshots:
        for index, screenshot in enumerate(screenshots, start=1):
            override = screenshot_overrides.get(index, {})
            title = choose_text(override.get("title"), label_from_filename(screenshot))
            description = choose_text(override.get("description_text"), screenshot_paragraph(title))
            parts.extend(
                [
                    "",
                    f"### 5.{index} {title}",
                    "",
                    f"![图5-{index} {title}](../../photos/{screenshot.name})",
                    "",
                    f"图5-{index} {title}",
                    "",
                    description,
                    "",
                    "\\newpage",
                ]
            )
    else:
        parts.extend(
            [
                "",
                "### 5.1 页面说明待补充",
                "",
                "当前尚未发现截图文件。先完成前端演示与 Playwright 截图，再补充本章节内容。",
            ]
        )

    return "\n".join(parts)


def build_final_docx(
    output_path: Path,
    system_name: str,
    short_name: str,
    frontend_language: str,
    backend_language: str,
    fullstack_languages: str,
    modules: list[str],
    screenshots: list[Path],
    manual_content: dict,
    author: str,
) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    for _ in range(5):
        document.add_paragraph("")

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(f"{system_name}[简称：{short_name}]V1.0")
    set_run_songti(title_run, 26, True)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run("产品说明书")
    set_run_songti(subtitle_run, 26, True)

    document.add_page_break()

    toc_title = document.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run("目  录")
    set_run_songti(toc_run, 22, True)
    toc_paragraph = document.add_paragraph()
    add_toc_field(toc_paragraph)

    document.add_page_break()

    add_heading_paragraph(document, "修订记录", 1)
    add_revision_table(document, author)

    document.add_page_break()

    add_heading_paragraph(document, "一、软件介绍", 1)
    for line in [
        f"软件名称：{system_name}",
        f"简称：{short_name}",
        "版本号：V1.0",
        "软件类别：应用软件",
        f"著作权人：{author}",
    ]:
        add_paragraph(document, line, first_line_indent_chars=2)

    document.add_page_break()

    add_heading_paragraph(document, "二、软件用途", 1)
    add_paragraph(
        document,
        choose_text(manual_content.get("purpose_text"), build_purpose_paragraph(system_name, modules)),
        first_line_indent_chars=2,
    )

    document.add_page_break()

    add_heading_paragraph(document, "三、软件功能", 1)
    add_paragraph(
        document,
        choose_text(manual_content.get("function_text"), build_function_paragraph(system_name, modules)),
        first_line_indent_chars=2,
    )

    document.add_page_break()

    add_heading_paragraph(document, "四、运行环境", 1)
    add_heading_paragraph(document, "4.1 硬件要求", 2)
    add_two_col_table(document, ("类型", "基本要求"), [("服务器端", "CPU 8核1.60GHz，内存8G，硬盘剩余空间10G")])
    document.add_paragraph("")
    add_heading_paragraph(document, "4.2 软件环境", 2)
    add_two_col_table(
        document,
        ("名称", "基本环境"),
        [
            ("操作系统", "Windows 10 64位"),
            ("数据库软件", "MySQL8.0"),
            ("开发软件", "Opencode，Claude Code，Codex，IntelliJ IDEA，Navicat 16"),
            ("开发语言", fullstack_languages),
        ],
    )

    document.add_page_break()

    add_heading_paragraph(document, "4.3 软件开发环境", 2)
    add_paragraph(
        document,
        choose_text(
            manual_content.get("development_environment_text"),
            f"本软件分为前端页面和后端业务逻辑，其中前端页面使用{frontend_language}进行开发，后端业务逻辑使用{backend_language}进行开发，数据库使用MySQL8.0，数据库管理采用Navicat 16，整个系统使用IntelliJ IDEA环境进行开发。开发界面如图4-1所示。",
        ),
        first_line_indent_chars=2,
    )
    add_paragraph(document, "[在此处粘贴开发界面图片]", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(document, "图4-1 软件开发界面", align=WD_ALIGN_PARAGRAPH.CENTER)

    document.add_page_break()

    add_heading_paragraph(document, "五、软件使用", 1)

    screenshot_overrides = extract_screenshot_overrides(manual_content)
    if screenshots:
        for index, screenshot in enumerate(screenshots, start=1):
            if index > 1:
                document.add_page_break()
            override = screenshot_overrides.get(index, {})
            title = choose_text(override.get("title"), label_from_filename(screenshot))
            add_heading_paragraph(document, f"5.{index} {title}", 2)
            maybe_add_picture(document, screenshot, choose_text(override.get("caption"), f"图5-{index} {title}"))
            add_paragraph(
                document,
                choose_text(override.get("description_text"), screenshot_paragraph(title)),
                first_line_indent_chars=2,
            )
    else:
        add_heading_paragraph(document, "5.1 页面说明待补充", 2)
        add_paragraph(document, "当前尚未发现截图文件。先完成前端演示与 Playwright 截图，再补充本章节内容。", first_line_indent_chars=2)

    document.save(output_path)


def run_pandoc(markdown_path: Path, output_path: Path) -> None:
    if not PANDOC_PATH.exists():
        return
    subprocess.run(
        [str(PANDOC_PATH), str(markdown_path), "-o", str(output_path)],
        check=True,
        cwd=str(markdown_path.parent),
    )


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Build a formatted manual DOCX using python-docx and Pandoc.",
    )
    parser.add_argument("--root", default=".", help="Workspace root")
    parser.add_argument("--system-name", required=True, help="Display name of the system")
    parser.add_argument("--author", default="孔祥鑫", help="Author name for revision history")
    parser.add_argument("--short-name", default="", help="Optional short product name")
    args = parser.parse_args()

    root = Path(args.root).resolve()
    system_name = display_system_name(args.system_name.strip())
    short_name = args.short_name.strip() or derive_short_name(system_name)
    manifest = load_manifest(root, args.system_name.strip())
    safe_name = manifest.get("system_folder") or safe_path_name(args.system_name.strip())
    system_dir = root / safe_name
    docs_dir = system_dir / "docs"
    template_dir = docs_dir / "Template"
    photos_dir = system_dir / "photos"
    docs_dir.mkdir(parents=True, exist_ok=True)
    template_dir.mkdir(parents=True, exist_ok=True)

    screenshots = sorted(
        [
            path
            for path in photos_dir.glob("*")
            if path.is_file() and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}
        ]
    )

    frontend_language = infer_frontend_language(system_dir)
    backend_language = infer_backend_language(system_dir)
    fullstack_languages = infer_fullstack_languages(system_dir)
    modules = infer_modules(system_dir)
    manual_content = load_manual_content(system_dir)

    errors, warnings = validate_manual_content(manual_content, screenshots)
    for warning in warnings:
        print(f"WARNING: {warning}")
    if errors:
        for error in errors:
            print(f"ERROR: {error}")
        print("Manual DOCX generation aborted because agent-written content is missing or still templated.")
        return 1

    markdown_draft = render_markdown_draft(
        system_name=system_name,
        short_name=short_name,
        frontend_language=frontend_language,
        backend_language=backend_language,
        fullstack_languages=fullstack_languages,
        modules=modules,
        screenshots=screenshots,
        manual_content=manual_content,
    )

    markdown_path = template_dir / f"{safe_name}-manual-draft.md"
    pandoc_docx_path = template_dir / f"{safe_name}-manual-pandoc-draft.docx"
    final_docx_path = docs_dir / f"{system_name}-系统说明书.docx"

    markdown_path.write_text(markdown_draft, encoding="utf-8")

    pandoc_status = "skipped"
    try:
        run_pandoc(markdown_path, pandoc_docx_path)
        if pandoc_docx_path.exists():
            pandoc_status = "created"
    except Exception as exc:
        pandoc_status = f"failed: {exc}"

    build_final_docx(
        output_path=final_docx_path,
        system_name=system_name,
        short_name=short_name,
        frontend_language=frontend_language,
        backend_language=backend_language,
        fullstack_languages=fullstack_languages,
        modules=modules,
        screenshots=screenshots,
        manual_content=manual_content,
        author=args.author,
    )

    print(f"Modules: {modules}")
    print(f"Short name: {short_name}")
    print(f"Markdown draft: {markdown_path}")
    print(f"Pandoc draft: {pandoc_docx_path}")
    print(f"Pandoc status: {pandoc_status}")
    print(f"Final DOCX: {final_docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
