#!/usr/bin/env python3
"""
Build the cooperation development agreement DOCX for a ruanzhu delivery pack.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from document_defaults import load_agreement_party_a

FONT_NAME = "宋体"
TITLE_SIZE_PT = 26
BODY_SIZE_PT = 12
PARTY_LABELS = ["甲", "乙"]
DEFAULT_AGREEMENT_DATE = "2026年4月15日"


def safe_path_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*\x00-\x1f]+', "-", value.strip())
    cleaned = re.sub(r"\s+", "-", cleaned)
    cleaned = re.sub(r"-{2,}", "-", cleaned).strip(" .-")
    return cleaned or "system-app"


def display_system_name(system_name: str) -> str:
    name = system_name.strip()
    return name if name.endswith("系统") else f"{name}系统"


def project_software_name(system_name: str) -> str:
    return display_system_name(system_name)


def load_manifest(root: Path, system_name: str) -> dict:
    direct_manifest = root / safe_path_name(system_name) / "docs" / "Template" / "delivery-manifest.json"
    if direct_manifest.exists():
        return json.loads(direct_manifest.read_text(encoding="utf-8"))

    direct_matches = list(root.glob("*/docs/Template/delivery-manifest.json"))
    if len(direct_matches) == 1:
        return json.loads(direct_matches[0].read_text(encoding="utf-8"))

    legacy_manifest = root / "outputs" / "Template" / "delivery-manifest.json"
    if legacy_manifest.exists():
        return json.loads(legacy_manifest.read_text(encoding="utf-8"))

    return {}


def resolve_system_dir(root: Path, system_name: str) -> Path:
    manifest = load_manifest(root, system_name)
    safe_name = manifest.get("system_folder") or safe_path_name(system_name)
    return root / safe_name


def parse_party_count(raw: str) -> int:
    text = str(raw).strip()
    digit_match = re.search(r"\d+", text)
    if digit_match:
        return int(digit_match.group(0))

    simple_map = {
        "二": 2,
        "两": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
        "十": 10,
    }
    for char, number in simple_map.items():
        if char in text:
            return number
    raise argparse.ArgumentTypeError(f"Invalid party count: {raw}")


def party_count_noun(count: int) -> str:
    if count != 2:
        raise ValueError("合作开发协议模板固定为甲乙双方，不再生成丙方及以上。")
    return "双方"


def default_agreement_date() -> str:
    return DEFAULT_AGREEMENT_DATE


def build_agreement_lines(
    system_name: str,
    party_count: int,
    agreement_date: str,
    party_a_name: str,
    party_a_id: str,
) -> list[str]:
    if party_count != 2:
        raise ValueError("合作开发协议模板固定为甲乙双方，不再生成丙方及以上。")
    labels = PARTY_LABELS[:party_count]
    party_phrase = "甲、乙双方"
    software_name = project_software_name(system_name)
    copies_text = "三"

    lines = []
    for index, label in enumerate(labels):
        if index == 0:
            lines.append(f"{label}方：{party_a_name}")
            lines.append(f"身份证号：{party_a_id}")
        else:
            lines.append(f"{label}方：")
            lines.append("身份证号：")

    lines.extend(
        [
            "鉴于，甲乙双方均为计算机软件专业开发人员，能够进行创造性的软件开发活动。并且，甲乙双方有意愿共同从事软件的开发工作。为了规范双方的权利义务，在《中华人民共和国合同法》及其他相关法规政策的原则指导下，订立本协议书，双方共同遵守：",
            "第一条、合作宗旨",
            f"{party_phrase}于{agreement_date}共同约定完成{software_name}的开发工作，并共同享有开发成果而合作。",
            "第二条、合作项目和范围",
            f"甲乙双方共同开发{software_name}，合作范围包括软件的代码编写、调试、测试等开发工作。",
            "第三条、合作方式",
            "1、甲乙双方按照软件编程工作的正常分工进行编写，任何一方不得随意更改软件的重大功能和事项，以免对对方造成履约困难。",
            "2、甲乙双方应坚持勤勉努力诚实信用的原则，进行双方分别负责的软件的编程工作，并考虑到双方软件的兼容和接合。如一方发生特殊技术困难，对方有义务为其提供合理适当的技术帮助。",
            "第四条、知识产权",
            "1、双方编写的软件源代码、技术文档及汇编而成的程序本身，其著作权均由甲乙双方共同享有。",
            "2、甲乙双方在编写软件的过程中，不得有侵犯他人知识产权的行为，否则，应对外承担全部侵权责任。",
            "第五条、协议变更",
            "1、经甲乙双方协商同意，本协议可以作相应变更。",
            "2、任何一方未经与对方协商，擅自变更本协议条款或者将本协议权利义务转让他人，均为无效。",
            "第六条、禁止行为",
            "1、未经双方一致同意，禁止任何一方私自以团体名义进行业务活动；如其业务获得利益归甲乙双方共有，造成损失按实际损失赔偿。",
            "2、禁止任何一方泄露本协议所涉及的相关商业秘密。",
            "第七条、合作的终止",
            "合作开发活动因以下事由之一得终止∶",
            "1、甲乙双方同意终止合作关系。",
            "2、合作项目因技术原因，根本不能完成。",
            "3、合作项目违反法律被撤销。",
            "第八条、违约责任",
            "1、在合作期内，项目合作双方中任一方未经对方协商认可擅自退出该合作项目，违约方同时赔偿被侵害方的投入损失及其他合作期内应得收益。并且必须遵守技术、市场保密条款，两年内不得在当地使用或经营本项目的同类技术内容及客户资源。否则守约方有权追究违约方的一切经济法律责任。",
            "2、在合作期内因战争、灾害、疾病等不可抗力因素导致项目合作解散或合作期满双方不再合作，该项目技术内容归甲乙双方所有。",
            "3、甲乙双方如有一方违反本协议，则另一方有权取消与违约方的合作并追究违约方的一切经济法律责任。",
            "第九条、纠纷的解决",
            "甲乙双方之间如发生纠纷，应共同协商，本着有利于事业发展的原则予以解决。如协商不成，可以诉诸法院。",
            "第十条、本协议如有未尽事宜，应由甲乙双方共同讨论补充或修改。补充和修改的内容与本协议具有同等效力。",
            f"第十一条、本协议一式{copies_text}份，{party_phrase}各执一份，交国家版权局备案一份。",
        ]
    )
    return lines


def ensure_rfonts(run, font_name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, size_pt: float) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    ensure_rfonts(run, FONT_NAME)


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, TITLE_SIZE_PT)


def add_body_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(BODY_SIZE_PT * 2)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, BODY_SIZE_PT)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)


def write_markdown_draft(path: Path, lines: list[str]) -> None:
    text = "# 合作开发协议\n\n" + "\n\n".join(lines) + "\n"
    path.write_text(text, encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate <system-name>合作开发协议.docx with fixed cooperation-agreement wording.",
    )
    parser.add_argument("--root", default=".", help="Workspace root")
    parser.add_argument("--system-name", required=True, help="Display name of the system")
    parser.add_argument("--party-count", type=parse_party_count, default=2, help="Number of agreement parties; fixed to 2")
    parser.add_argument(
        "--agreement-date",
        default=None,
        help="Date text used in Article 1; defaults to 2026年4月15日",
    )
    args = parser.parse_args()

    if args.party_count != 2:
        raise ValueError("party-count must be 2; the agreement template is fixed to 甲乙双方.")

    root = Path(args.root).resolve()
    system_display_name = display_system_name(args.system_name.strip())
    manifest = load_manifest(root, args.system_name.strip())
    system_dir = resolve_system_dir(root, args.system_name.strip())
    docs_dir = system_dir / "docs"
    template_dir = docs_dir / "Template"
    docs_dir.mkdir(parents=True, exist_ok=True)
    template_dir.mkdir(parents=True, exist_ok=True)

    preferred_date = manifest.get("delivery_preferences", {}).get("agreement_date", "")
    agreement_date = args.agreement_date.strip() if args.agreement_date else preferred_date or default_agreement_date()
    party_a_name, party_a_id = load_agreement_party_a()
    lines = build_agreement_lines(
        system_display_name,
        args.party_count,
        agreement_date,
        party_a_name,
        party_a_id,
    )

    document = Document()
    configure_document(document)
    document.core_properties.title = "合作开发协议"
    add_title(document, "合作开发协议")
    for line in lines:
        add_body_line(document, line)

    draft_path = template_dir / f"{safe_path_name(system_display_name)}-cooperation-agreement-draft.md"
    final_docx = docs_dir / f"{system_display_name}合作开发协议.docx"
    write_markdown_draft(draft_path, lines)
    document.save(final_docx)

    print(f"Agreement draft: {draft_path}")
    print(f"Agreement DOCX: {final_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
